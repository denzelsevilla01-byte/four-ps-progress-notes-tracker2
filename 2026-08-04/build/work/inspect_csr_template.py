from docx import Document
from pathlib import Path
import hashlib, json, sys
p=Path(r"C:\Users\WINDOWS\Desktop\PANTAWID FILES\MOO-SAN FABIAN\MOO-SAN FABIAN\CASE FOLDERS\CSR\CSR-ACTIVE\AMBALANGAN-DALIN\SCSR_Mariza E. Salango.docx")
d=Document(p)
print(json.dumps({"sha256":hashlib.sha256(p.read_bytes()).hexdigest(),"sections":len(d.sections),"paragraphs":len(d.paragraphs),"tables":len(d.tables)},indent=2))
for i,para in enumerate(d.paragraphs):
    if para.text.strip(): print(f"P{i} [{para.style.name}] {para.text}")
for ti,t in enumerate(d.tables):
    print(f"TABLE {ti} {len(t.rows)}x{len(t.columns)}")
    for ri,row in enumerate(t.rows):
        print(f" R{ri}: "+" || ".join(c.text.replace("\n"," / ") for c in row.cells))
for si,s in enumerate(d.sections):
    print(f"SECTION {si} {s.page_width}x{s.page_height} margins={s.top_margin},{s.right_margin},{s.bottom_margin},{s.left_margin}")
    print(" HEADER:"," | ".join(x.text for x in s.header.paragraphs if x.text.strip()))
    print(" FOOTER:"," | ".join(x.text for x in s.footer.paragraphs if x.text.strip()))
